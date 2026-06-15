"""
Run this script to generate the Word document:
    pip install python-docx
    python generate_doc.py
Output: docs/Healthcare_Pipeline_Documentation.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)
IMG  = os.path.join(ROOT, "images")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Helpers ───────────────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"),   "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"),  "F2F2F2")
    p._p.pPr.append(shading) if p._p.pPr is not None else None
    return p

def img(path, width=Inches(5.5), caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].font.size   = Pt(9)
            cp.runs[0].font.italic = True
            cp.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    else:
        body(f"[Image not found: {path}]")

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"),   "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"),  "1F497D")
        hdr[i]._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = w
    doc.add_paragraph()
    return t

def divider():
    doc.add_paragraph("─" * 80)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
title = doc.add_heading("Real-Time Healthcare Claims &\nEHR Data Engineering Pipeline on AWS", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.size = Pt(24)

doc.add_paragraph()
sub = doc.add_paragraph("End-to-End Pipeline: Lambda → S3 → Glue → Snowflake → Streamlit")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(13)
sub.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
meta = doc.add_paragraph("Authors: Thrivikram Kotharu, Arkala\nDate: June 2026\nStack: AWS Lambda · AWS Glue · Amazon S3 · Snowflake · Streamlit")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

heading("1. Project Overview")
body("This project builds a real-time healthcare data engineering pipeline that:")
bullet("Generates synthetic CMS Claims (Medicare/Medicaid style) and FHIR R4 EHR data")
bullet("Ingests raw data into Amazon S3 via AWS Lambda functions")
bullet("Transforms and cleans the data using AWS Glue PySpark ETL jobs")
bullet("Loads clean Parquet data into Snowflake automatically via Snowpipe")
bullet("Provides an interactive Streamlit analytics dashboard for exploring the data")
body("The pipeline is fully event-driven: a single API call triggers data generation, which cascades through S3 → Glue → Snowflake with no manual intervention.")

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

heading("2. Architecture")
body("The pipeline follows a layered, event-driven architecture:")

arch = """USER  →  API Gateway  →  Lambda: api-trigger
                                      │
                     ┌────────────────┴────────────────┐
                     ▼                                   ▼
          Lambda: cms-claims-generator       Lambda: fhir-ehr-generator
                     │                                   │
                     ▼                                   ▼
          S3: healthcare-cms-claims-raw      S3: healthcare-ehr-raw
                     │                                   │
                     ▼                                   ▼
          Glue: cms-claims-etl              Glue: ehr-fhir-etl
                     │                                   │
                     ▼                                   ▼
          S3: cms-claims-clean             S3: ehr-clean
                     │                                   │
                     └──────────────┬────────────────────┘
                                    ▼
                            Snowflake Snowpipe
                                    │
                            Snowflake Tables
                                    │
                          Streamlit Dashboard"""
code(arch)

# ══════════════════════════════════════════════════════════════════════════════
# 3. DATA SOURCES & SCHEMA
# ══════════════════════════════════════════════════════════════════════════════

heading("3. Data Sources & Schema")

heading("3.1  CMS Claims Data (Synthetic DE-SynPUF)", level=2)
body("The CMS DE-SynPUF (Synthetic Public Use File) schema mirrors real Medicare claims structure. Seven entity types are generated per batch:")

table(
    ["Entity", "Description", "Key Fields"],
    [
        ["beneficiary_summary", "Patient demographics + chronic conditions", "DESYNPUF_ID, birth date, gender, race, 9 chronic flags"],
        ["providers",           "Provider directory",                         "PROVIDER_ID, NPI, specialty, state"],
        ["inpatient_claims",    "Hospital stays",                             "CLM_ID, admit/discharge dates, DRG code, payment"],
        ["outpatient_claims",   "Outpatient visits",                          "CLM_ID, service dates, HCPCS codes, payment"],
        ["carrier_claims",      "Physician/carrier claims",                   "CLM_ID, performing NPI, service date, payment"],
        ["prescription_drug_events", "Pharmacy fills",                       "PDE_ID, drug code, generic/brand name, cost"],
        ["diagnosis_codes",     "ICD-10 reference codes",                    "ICD10 code + description"],
    ],
    col_widths=[Inches(1.8), Inches(2.0), Inches(2.9)]
)

body("Medical coding standards used: ICD-10 (diagnoses), HCPCS (procedures), DRG (inpatient classification), NDC (drug codes).")

heading("3.2  FHIR R4 EHR Data (Synthetic Bundles)", level=2)
body("FHIR (Fast Healthcare Interoperability Resources) R4 is the modern standard for healthcare data exchange. Each patient generates a Bundle containing 15 linked resources:")

table(
    ["FHIR Resource", "Description", "Key Fields"],
    [
        ["Patient",              "Patient demographics",        "id, name, gender, birthDate, address, race, ethnicity"],
        ["Practitioner",         "Healthcare provider",         "id, NPI, name, qualification/specialty"],
        ["Encounter",            "Clinical visit",              "id, class (AMB/EMER/IMP), period, total cost"],
        ["Condition",            "Diagnoses",                   "id, ICD-10, SNOMED code, clinical status, chronic flag"],
        ["Observation",          "Lab results & vitals",        "id, LOINC code, value, unit, reference range, abnormal flag"],
        ["MedicationRequest",    "Prescriptions",               "id, RxNorm code, drug name, dosage, chronic flag"],
        ["Procedure",            "Performed procedures",        "id, CPT code, SNOMED code, cost"],
        ["AllergyIntolerance",   "Allergies",                   "id, substance, reaction, severity"],
        ["Immunization",         "Vaccines",                    "id, CVX code, vaccine name, occurrence date"],
    ],
    col_widths=[Inches(1.6), Inches(1.8), Inches(3.2)]
)

body("Medical coding standards used: ICD-10, SNOMED-CT, LOINC (lab tests), RxNorm (medications), CVX (vaccines), CPT (procedures).")

# ══════════════════════════════════════════════════════════════════════════════
# 4. AWS INFRASTRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

heading("4. AWS Infrastructure Setup")
body("All AWS services were created manually via the AWS Console. Region used throughout: us-east-2 (Ohio).")

heading("4.1  IAM Roles", level=2)
table(
    ["Role", "Used By", "Key Permissions"],
    [
        ["healthcare-lambda-role", "All Lambda functions",  "S3 read/write, Lambda invoke, CloudWatch logs"],
        ["healthcare-glue-role",   "Glue ETL jobs",         "S3 read/write (raw + clean buckets), CloudWatch logs"],
        ["snowflake-s3-role",      "Snowflake integration", "S3 GetObject, ListBucket on clean buckets"],
    ],
    col_widths=[Inches(2.0), Inches(1.8), Inches(2.9)]
)

heading("4.2  S3 Buckets", level=2)
table(
    ["Bucket", "Purpose"],
    [
        ["healthcare-cms-claims-raw",   "Raw CMS CSV files + Glue ETL scripts"],
        ["healthcare-ehr-raw",          "Raw FHIR R4 JSON bundles"],
        ["healthcare-cms-claims-clean", "Cleaned CMS Parquet files (Snowpipe source)"],
        ["healthcare-ehr-clean",        "Cleaned EHR Parquet files (Snowpipe source)"],
    ],
    col_widths=[Inches(3.0), Inches(3.6)]
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. LAMBDA FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

heading("5. Data Ingestion — Lambda Functions")
body("Three Lambda functions handle data generation and orchestration. All use Python 3.11 and the handler.handler entry point.")

heading("5.1  cms-claims-generator", level=2)
body("Generates synthetic CMS Claims CSVs and uploads to S3 under date-partitioned prefixes.")
bullet("Trigger: Async invocation from api-trigger Lambda")
bullet("Output: 7 CSV files per invocation")
bullet("Key env vars: RAW_CMS_BUCKET, RECORDS_PER_BATCH (default 500)")
body("S3 key pattern:")
code("year=2026/month=06/day=15/beneficiary_summary_<uuid>.csv\nyear=2026/month=06/day=15/inpatient_claims_<uuid>.csv\nyear=2026/month=06/day=15/carrier_claims_<uuid>.csv")

heading("5.2  fhir-ehr-generator", level=2)
body("Generates synthetic FHIR R4 Bundle JSON files and uploads to S3.")
bullet("Trigger: Async invocation from api-trigger Lambda")
bullet("Output: One JSON Bundle per patient under fhir/")
bullet("Key env vars: RAW_EHR_BUCKET, PATIENTS_PER_BATCH (default 200)")
bullet("Structure: Each file is a FHIR Bundle containing 15 linked resources for one patient")

heading("5.3  api-trigger", level=2)
body("Entry point exposed via API Gateway. Asynchronously invokes both generators in parallel.")
bullet("Trigger: HTTP POST via API Gateway")
bullet("InvocationType=Event — fire and forget, both generators run in parallel")
bullet("Returns immediately with a batchId while generation continues in the background")

img(os.path.join(IMG, "AWS", "lambda function.jpeg"), caption="Figure 1 — Lambda Functions in AWS Console")

# ══════════════════════════════════════════════════════════════════════════════
# 6. API GATEWAY
# ══════════════════════════════════════════════════════════════════════════════

heading("6. API Gateway — Trigger Endpoint")
body("An HTTP API (API Gateway v2) exposes a single POST endpoint that kicks off the entire pipeline.")

body("Endpoint URL:")
code("POST https://u0zg8lmc1c.execute-api.us-east-2.amazonaws.com/dev/trigger")

body("Request body:")
code('{\n  "records_per_batch": 500,\n  "patients_per_batch": 200\n}')

body("Response:")
code('{\n  "message": "Data generation triggered successfully.",\n  "batchId": "2c8b7229-03c8-4748-b8c0-6eda01032cd0"\n}')

body("How to call it (PowerShell):")
code('Invoke-RestMethod \\\n  -Uri "https://u0zg8lmc1c.execute-api.us-east-2.amazonaws.com/dev/trigger" \\\n  -Method POST \\\n  -ContentType "application/json" \\\n  -Body \'{"records_per_batch": 500, "patients_per_batch": 200}\'')

img(os.path.join(IMG, "AWS", "api gateway.jpeg"), caption="Figure 2 — API Gateway Configuration")
img(os.path.join(IMG, "AWS", "Api gateway image representation.jpeg"), caption="Figure 3 — API Gateway Architecture")

# ══════════════════════════════════════════════════════════════════════════════
# 7. RAW DATA STORAGE — S3
# ══════════════════════════════════════════════════════════════════════════════

heading("7. Raw Data Storage — Amazon S3")
body("Raw data lands in two S3 buckets with date-partitioned prefixes, allowing Glue to process data by date range.")

heading("7.1  CMS Raw Bucket Structure", level=2)
code("healthcare-cms-claims-raw/\n├── glue-scripts/\n│   ├── cms_claims_etl.py\n│   └── ehr_fhir_etl.py\n└── year=2026/\n    └── month=06/\n        └── day=15/\n            ├── beneficiary_summary_<uuid>.csv\n            ├── providers_<uuid>.csv\n            ├── inpatient_claims_<uuid>.csv\n            ├── outpatient_claims_<uuid>.csv\n            ├── carrier_claims_<uuid>.csv\n            ├── prescription_drug_events_<uuid>.csv\n            └── diagnosis_codes_<uuid>.csv")

heading("7.2  EHR Raw Bucket Structure", level=2)
code("healthcare-ehr-raw/\n└── fhir/\n    ├── patient_<uuid>.json\n    ├── patient_<uuid>.json\n    └── ...")

img(os.path.join(IMG, "AWS", "S3.jpeg"),               caption="Figure 4 — S3 Buckets Overview")
img(os.path.join(IMG, "AWS", "S3-1.jpeg"),             caption="Figure 5 — S3 Raw Bucket Contents")
img(os.path.join(IMG, "AWS", "s3 folder names.jpeg"),  caption="Figure 6 — S3 Date-Partitioned Folder Structure")
img(os.path.join(IMG, "AWS", "s3 folder names 2.jpeg"),caption="Figure 7 — S3 Entity File Names")

# ══════════════════════════════════════════════════════════════════════════════
# 8. AWS GLUE ETL
# ══════════════════════════════════════════════════════════════════════════════

heading("8. ETL — AWS Glue")
body("Two AWS Glue 4.0 PySpark ETL jobs transform raw data into analytics-ready Parquet format.")

heading("8.1  CMS Claims ETL (cms-claims-etl)", level=2)
body("Script: src/glue/cms_claims_etl.py | Runtime: Glue 4.0, PySpark 3.3, Python 3.10")
body("Job Parameters:")
code("--RAW_CMS_BUCKET    healthcare-cms-claims-raw\n--CLEAN_CMS_BUCKET  healthcare-cms-claims-clean")

body("Nine transformation groups are applied in sequence:")
table(
    ["Group", "What It Does"],
    [
        ["1. Column rename",    "Maps raw CMS column names (e.g. DESYNPUF_ID) to clean names (patient_id)"],
        ["2. Type casting",     "Converts date strings (yyyyMMdd) to DATE, amounts to DECIMAL(12,2)"],
        ["3. Code decoding",    "Maps numeric gender/race codes to labels (1→Male, 1→White)"],
        ["4. Null handling",    "Fills amount nulls with 0.0, flags missing birth dates, sets is_deceased"],
        ["5. Deduplication",    "dropDuplicates on business keys (patient_id, claim_id, etc.)"],
        ["6. Date derivation",  "Extracts year, quarter, month, week_of_year, date_key from date columns"],
        ["7. Financial cleanup","Clips negative amounts to 0, rounds to 2dp, flags is_high_cost at 99th pct"],
        ["8. Surrogate keys",   "SHA-256 hash of business key → patient_key, claim_key, etc."],
        ["9. Risk flags",       "Sums 9 chronic condition flags → risk_score, risk_tier (Low/Medium/High)"],
    ],
    col_widths=[Inches(1.8), Inches(4.8)]
)

body("Output tables written as Parquet:")
table(
    ["Parquet Prefix", "Description", "Sample Rows"],
    [
        ["dim_patient/",        "Patient dimension",               "3,000"],
        ["dim_provider/",       "Provider dimension",              "300"],
        ["fact_claims/",        "Inpatient + outpatient claims",   "6,000"],
        ["fact_carrier_claims/","Carrier/physician claims",        "3,000"],
        ["fact_prescriptions/", "Drug events",                     "3,000"],
    ],
    col_widths=[Inches(1.8), Inches(2.5), Inches(1.5)]
)

heading("8.2  EHR FHIR ETL (ehr-fhir-etl)", level=2)
body("Script: src/glue/ehr_fhir_etl.py | Runtime: Glue 4.0, PySpark 3.3, Python 3.10")
body("Job Parameters:")
code("--RAW_EHR_BUCKET    healthcare-ehr-raw\n--CLEAN_EHR_BUCKET  healthcare-ehr-clean")

body("Key challenge — FHIR Schema Conflicts:")
body("FHIR bundles contain mixed resource types in a single JSON array. Spark infers a merged schema, causing three type conflicts that were resolved:")
bullet("name field conflict: Organization.name is a plain string; Patient.name is an array of structs. Fixed using from_json() with an explicit schema after filtering by resourceType.")
bullet("valueMoney vs valueDecimal: The generator uses valueMoney:{value,currency} for costs. Fixed by changing to extension[0]['valueMoney']['value'].")
bullet("category field conflict: AllergyIntolerance.category is ['food'] (string array); Observation.category is [{coding:[...]}] (struct array). Fixed using get_json_object() instead of struct access.")

body("Output tables:")
table(
    ["Parquet Prefix", "Description", "Sample Rows"],
    [
        ["dim_patient/",      "EHR patient dimension",      "532"],
        ["dim_provider/",     "Practitioner dimension",     "~50"],
        ["fact_encounters/",  "Clinical visits",            "532"],
        ["fact_conditions/",  "Diagnoses",                  "0*"],
        ["fact_observations/","Lab results & vitals",       "1,596"],
        ["fact_medications/", "Medication requests",        "1,064"],
    ],
    col_widths=[Inches(1.8), Inches(2.5), Inches(1.5)]
)
body("* fact_conditions has 0 rows in this sample batch — varies by batch size.")

img(os.path.join(IMG, "AWS", "Glue runs.jpeg"),     caption="Figure 8 — AWS Glue Job Runs")
img(os.path.join(IMG, "AWS", "cms glue runs.jpeg"), caption="Figure 9 — CMS Glue Job Run History")
img(os.path.join(IMG, "AWS", "glue logs.jpeg"),     caption="Figure 10 — Glue Job CloudWatch Logs")

# ══════════════════════════════════════════════════════════════════════════════
# 9. CLEAN S3
# ══════════════════════════════════════════════════════════════════════════════

heading("9. Clean Data Storage — S3 Clean Buckets")
body("After Glue completes, cleaned Parquet files land in the clean S3 buckets. Snowpipe monitors these buckets and loads new files automatically.")

body("Why Parquet?")
bullet("Columnar format — fast for analytical queries (reads only needed columns)")
bullet("Snappy compression — ~75% smaller than equivalent CSV")
bullet("Type-preserving — dates, decimals, and booleans stored correctly without parsing")
bullet("Native support in both Apache Spark and Snowflake")

# ══════════════════════════════════════════════════════════════════════════════
# 10. SNOWFLAKE
# ══════════════════════════════════════════════════════════════════════════════

heading("10. Snowflake — Data Warehouse")

heading("10.1  Account Details", level=2)
table(
    ["Field", "Value"],
    [
        ["Account Identifier", "ABCIDUX-GP18792"],
        ["Organization",       "ABCIDUX"],
        ["Account Name",       "GP18792"],
        ["Login Name",         "ARKALA"],
        ["Warehouse",          "COMPUTE_WH (X-Small)"],
        ["Database",           "healthcare"],
        ["Schema",             "raw"],
    ],
    col_widths=[Inches(2.0), Inches(4.5)]
)

heading("10.2  Tables Loaded", level=2)
body("CMS Tables:")
table(
    ["Table", "Source", "Rows"],
    [
        ["dim_patient_cms",   "beneficiary_summary",      "3,000"],
        ["dim_provider_cms",  "providers",                "300"],
        ["fact_claims",       "inpatient + outpatient",   "6,000"],
        ["fact_carrier_claims","carrier_claims",          "3,000"],
        ["fact_prescriptions","prescription_drug_events", "3,000"],
    ],
    col_widths=[Inches(2.0), Inches(2.5), Inches(1.5)]
)
body("EHR Tables:")
table(
    ["Table", "Source", "Rows"],
    [
        ["dim_patient_ehr",   "FHIR Patient",          "532"],
        ["dim_provider_ehr",  "FHIR Practitioner",     "~50"],
        ["fact_encounters",   "FHIR Encounter",        "532"],
        ["fact_conditions",   "FHIR Condition",        "0"],
        ["fact_observations", "FHIR Observation",      "1,596"],
        ["fact_medications",  "FHIR MedicationRequest","1,064"],
    ],
    col_widths=[Inches(2.0), Inches(2.5), Inches(1.5)]
)

heading("10.3  Storage Integration", level=2)
body("A Snowflake Storage Integration connects Snowflake to both clean S3 buckets using an IAM role trust relationship — no AWS credentials stored in Snowflake.")
code("CREATE STORAGE INTEGRATION s3_healthcare\n    TYPE                      = EXTERNAL_STAGE\n    STORAGE_PROVIDER          = 'S3'\n    ENABLED                   = TRUE\n    STORAGE_AWS_ROLE_ARN      = 'arn:aws:iam::109676855909:role/snowflake-s3-role'\n    STORAGE_ALLOWED_LOCATIONS = (\n        's3://healthcare-cms-claims-clean/',\n        's3://healthcare-ehr-clean/'\n    );")

img(os.path.join(IMG, "snowflake", "snowflake.png"), caption="Figure 11 — Snowflake Tables and Data")

# ══════════════════════════════════════════════════════════════════════════════
# 11. SNOWPIPE
# ══════════════════════════════════════════════════════════════════════════════

heading("11. Snowpipe — Auto-Ingest")
body("Snowpipe automatically loads new Parquet files into Snowflake tables as soon as they appear in S3 — no manual COPY INTO required.")

heading("11.1  How It Works", level=2)
code("Glue writes Parquet to S3 clean bucket\n        ↓\nS3 fires ObjectCreated event → SQS queue (owned by Snowflake)\n        ↓  (~60 seconds latency)\nSnowpipe reads SQS message, loads new files into target table\n        ↓\nData available to query in Snowflake")

heading("11.2  Pipes Created", level=2)
body("CMS Pipes (5):")
table(
    ["Pipe Name", "Source Prefix", "Target Table"],
    [
        ["pipe_dim_patient_cms",   "cms_clean_stage/dim_patient/",        "dim_patient_cms"],
        ["pipe_dim_provider_cms",  "cms_clean_stage/dim_provider/",       "dim_provider_cms"],
        ["pipe_fact_claims",       "cms_clean_stage/fact_claims/",        "fact_claims"],
        ["pipe_fact_carrier",      "cms_clean_stage/fact_carrier_claims/","fact_carrier_claims"],
        ["pipe_fact_prescriptions","cms_clean_stage/fact_prescriptions/", "fact_prescriptions"],
    ],
    col_widths=[Inches(2.2), Inches(2.8), Inches(1.7)]
)
body("EHR Pipes (6):")
table(
    ["Pipe Name", "Source Prefix", "Target Table"],
    [
        ["pipe_dim_patient_ehr",  "ehr_clean_stage/dim_patient/",    "dim_patient_ehr"],
        ["pipe_dim_provider_ehr", "ehr_clean_stage/dim_provider/",   "dim_provider_ehr"],
        ["pipe_fact_encounters",  "ehr_clean_stage/fact_encounters/","fact_encounters"],
        ["pipe_fact_conditions",  "ehr_clean_stage/fact_conditions/","fact_conditions"],
        ["pipe_fact_observations","ehr_clean_stage/fact_observations/","fact_observations"],
        ["pipe_fact_medications", "ehr_clean_stage/fact_medications/","fact_medications"],
    ],
    col_widths=[Inches(2.2), Inches(2.8), Inches(1.7)]
)

heading("11.3  S3 Event Notifications", level=2)
body("Each clean S3 bucket has prefix-level ObjectCreated event notifications pointing to the Snowpipe SQS queue (ARN: arn:aws:sqs:us-east-2:109676855909:sf-snowpipe-...). All 11 pipes share the same SQS ARN — Snowflake routes events internally based on the stage path.")

body("Verification:")
code("SHOW PIPES;\n\nSELECT * FROM TABLE(INFORMATION_SCHEMA.COPY_HISTORY(\n    TABLE_NAME => 'FACT_CLAIMS',\n    START_TIME => DATEADD(HOUR, -2, CURRENT_TIMESTAMP)\n));")

# ══════════════════════════════════════════════════════════════════════════════
# 12. STREAMLIT DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════

heading("12. Streamlit Analytics Dashboard")
body("An interactive Streamlit dashboard connects directly to Snowflake using the snowflake-connector-python library and renders live charts from query results.")

heading("12.1  Setup", level=2)
code("cd streamlit_app\npython -m venv venv\nvenv\\Scripts\\activate\npip install -r requirements.txt\nstreamlit run app.py")

body(".env file (credentials — never committed to git):")
code("SNOWFLAKE_ACCOUNT=ABCIDUX-GP18792\nSNOWFLAKE_USER=Arkala\nSNOWFLAKE_WAREHOUSE=COMPUTE_WH\nSNOWFLAKE_DATABASE=healthcare\nSNOWFLAKE_SCHEMA=raw")

heading("12.2  Dashboard Tabs", level=2)

body("Overview Tab — KPIs + Patient Demographics:")
bullet("9 KPI metric tiles (patients, claims, spend, encounters, etc.)")
bullet("Patient Risk Distribution — pie chart derived from risk_score (Low/Medium/High)")
bullet("Patient Age Group — bar chart with average risk score as color scale")
bullet("Gender × Race Breakdown — grouped bar chart")

img(os.path.join(IMG, "streamlit dashboard", "streamlitdashboard1.png"), caption="Figure 12 — Streamlit Overview Tab")

body("CMS Claims Tab:")
bullet("Claims by Month — bar (count) + line (total payment) using YEAR/MONTH(admit_date)")
bullet("Payment Amount Buckets — claim count by payment range (< $500, $500–$2K, etc.)")
bullet("Top 10 Prescribed Drugs — horizontal bar colored by total cost")
bullet("High-Cost Patients Scatter — bubble chart: risk score vs total spend, sized by claim count")

img(os.path.join(IMG, "streamlit dashboard", "streamlitdashboard2.png"), caption="Figure 13 — CMS Claims Tab")

body("EHR Tab:")
bullet("Encounter Class Breakdown — pie (AMB/EMER/IMP) + avg cost bar chart")
bullet("Encounters by Month — bar + line using YEAR/MONTH(period_start)")
bullet("Top 10 Medications — horizontal bar colored by chronic medication count")
bullet("Chronic vs Acute Medication Split — donut pie chart")
bullet("Lab Tests by Volume — colored by abnormal percentage")
bullet("Abnormal Rate by Lab Test — percentage horizontal bar")

img(os.path.join(IMG, "streamlit dashboard", "streamlitdashboard3.png"), caption="Figure 14 — EHR Tab (Encounters & Medications)")
img(os.path.join(IMG, "streamlit dashboard", "streamlitdashboard4.png"), caption="Figure 15 — EHR Tab (Lab Observations)")

body("Diagnostics Tab:")
bullet("Row counts per table — bar chart showing how many rows Snowpipe has loaded")
bullet("Sample Data Viewer — dropdown to preview any table (5 rows)")
bullet("Null Detection — highlights columns with null values in the sample")

heading("12.3  Technical Note on Null Columns", level=2)
body("Spark's partitionBy() writes partition column values into the S3 folder path (e.g., fact_claims/year=2026/) and removes them from the Parquet file schema. Snowpipe reads raw Parquet files without path inference, so those columns (year, claim_type, visit_year, etc.) arrive as NULL in Snowflake. The dashboard works around this by deriving year/month directly from admit_date and period_start using YEAR() and MONTH() SQL functions.")

# ══════════════════════════════════════════════════════════════════════════════
# 13. CI/CD
# ══════════════════════════════════════════════════════════════════════════════

heading("13. CI/CD — GitHub Actions")
body("A GitHub Actions workflow (.github/workflows/deploy.yml) automatically deploys code on every push to main.")

heading("13.1  What It Deploys", level=2)
body("Job 1 — deploy-lambdas (runs in parallel with Job 2):")
bullet("Packages and deploys all 3 Lambda functions with aws lambda update-function-code")
bullet("Creates the function automatically if it does not exist yet")

body("Job 2 — deploy-glue:")
bullet("Uploads cms_claims_etl.py and ehr_fhir_etl.py to s3://healthcare-cms-claims-raw/glue-scripts/")
bullet("Updates the Glue job script location via aws glue update-job")

heading("13.2  Required Secrets & Variables", level=2)
table(
    ["Type", "Name", "Value"],
    [
        ["Secret",   "AWS_ACCESS_KEY_ID",       "IAM user access key"],
        ["Secret",   "AWS_SECRET_ACCESS_KEY",   "IAM user secret key"],
        ["Secret",   "LAMBDA_ROLE_ARN",          "arn:aws:iam::109676855909:role/healthcare-lambda-role"],
        ["Variable", "CMS_RAW_BUCKET",           "healthcare-cms-claims-raw"],
        ["Variable", "EHR_RAW_BUCKET",           "healthcare-ehr-raw"],
        ["Variable", "CMS_LAMBDA_NAME",          "cms-claims-generator"],
        ["Variable", "FHIR_LAMBDA_NAME",         "fhir-ehr-generator"],
        ["Variable", "TRIGGER_LAMBDA_NAME",      "api-trigger"],
        ["Variable", "CMS_GLUE_JOB",             "cms-claims-etl"],
        ["Variable", "EHR_GLUE_JOB",             "ehr-fhir-etl"],
    ],
    col_widths=[Inches(0.8), Inches(2.2), Inches(3.5)]
)

# ══════════════════════════════════════════════════════════════════════════════
# 14. END-TO-END FLOW
# ══════════════════════════════════════════════════════════════════════════════

heading("14. End-to-End Data Flow")
body("Complete journey of a single API trigger call:")

steps = [
    ("Step 1 — API Trigger",
     "User calls POST /trigger → API Gateway routes to api-trigger Lambda → returns batchId immediately"),
    ("Step 2 — Parallel Generation",
     "api-trigger async invokes both generators in parallel:\n  • cms-claims-generator: generates 500 beneficiaries + 6 entity files → uploads 7 CSVs to healthcare-cms-claims-raw/year=.../\n  • fhir-ehr-generator: generates 200 FHIR Bundles → uploads JSON to healthcare-ehr-raw/fhir/"),
    ("Step 3 — Glue ETL (CMS)",
     "cms-claims-etl reads all CSVs → applies 9 transformation groups → writes 5 Parquet tables to healthcare-cms-claims-clean/"),
    ("Step 4 — Glue ETL (EHR)",
     "ehr-fhir-etl reads FHIR JSON bundles → explodes Bundle.entry[] → flattens 9 resource types → writes Parquet to healthcare-ehr-clean/"),
    ("Step 5 — S3 Events",
     "S3 fires ObjectCreated events on clean buckets → routed to Snowflake SQS queues (~60 sec latency)"),
    ("Step 6 — Snowpipe Load",
     "Snowpipe reads SQS messages → COPY INTO target tables using MATCH_BY_COLUMN_NAME=CASE_INSENSITIVE"),
    ("Step 7 — Dashboard",
     "Streamlit queries Snowflake → renders interactive charts for CMS Claims, EHR, and diagnostics"),
]

for title_step, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title_step + ":  ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 15. FILE STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

heading("15. Project File Structure")
code("""Real-Time-Healthcare-Claims-EHR-Data-Engineering-Pipeline-on-AWS/
├── .github/
│   └── workflows/
│       └── deploy.yml                  ← CI/CD pipeline
├── src/
│   ├── lambdas/
│   │   ├── api_trigger/                ← Lambda 3: HTTP entry point
│   │   ├── cms_claims_generator/       ← Lambda 1: CMS data generator
│   │   └── fhir_generator/             ← Lambda 2: FHIR data generator
│   └── glue/
│       ├── cms_claims_etl.py           ← PySpark CMS ETL (9 transform groups)
│       └── ehr_fhir_etl.py             ← PySpark FHIR ETL (9 resource types)
├── streamlit_app/
│   ├── app.py                          ← Streamlit dashboard (4 tabs)
│   ├── queries.py                      ← All SQL queries
│   ├── requirements.txt
│   └── .env                            ← Snowflake credentials (gitignored)
├── snowflake/
│   ├── 01_setup.sql                    ← Database + schema
│   ├── 02_storage_integration.sql      ← S3 integration + IAM trust
│   ├── 03_stages.sql                   ← File format + S3 stages
│   ├── 04_tables_cms.sql               ← CMS table DDL
│   ├── 05_tables_ehr.sql               ← EHR table DDL
│   ├── 06_pipes_cms.sql                ← CMS Snowpipes
│   ├── 07_pipes_ehr.sql                ← EHR Snowpipes
│   └── 08_analytics_queries.sql        ← All dashboard queries
├── docs/
│   ├── pipeline_documentation.md       ← Markdown version
│   └── Healthcare_Pipeline_Documentation.docx ← This document
└── images/
    ├── AWS/                            ← AWS console screenshots
    ├── snowflake/                      ← Snowflake screenshots
    └── streamlit dashboard/            ← Dashboard screenshots""")

# ── Save ──────────────────────────────────────────────────────────────────────

out = os.path.join(BASE, "Healthcare_Pipeline_Documentation.docx")
doc.save(out)
print(f"Document saved: {out}")
